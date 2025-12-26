import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docxtpl import DocxTemplate  # 核心：用於模板填充
from io import BytesIO
import google.generativeai as genai
import json
import os
from datetime import datetime
import gspread
from google.oauth2.service_account import Credentials

# ==========================================
# 0. 頁面基本設定
# ==========================================
st.set_page_config(
    page_title="數位產業署政策規劃組行政秘書", 
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 1. 📊 本地用量記帳系統
# ==========================================
USAGE_LOG_FILE = "usage_log.json"

def load_usage_data():
    today_str = datetime.now().strftime("%Y-%m-%d")
    default_data = {"date": today_str, "stats": {}}
    if not os.path.exists(USAGE_LOG_FILE):
        return default_data
    try:
        with open(USAGE_LOG_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if data.get("date") != today_str:
                return default_data 
            return data
    except:
        return default_data

def save_usage_data(data):
    with open(USAGE_LOG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def update_usage_count(model_name, input_tokens, output_tokens):
    data = load_usage_data()
    if model_name not in data["stats"]:
        data["stats"][model_name] = {"count": 0, "total_tokens": 0}
    data["stats"][model_name]["count"] += 1
    data["stats"][model_name]["total_tokens"] += (input_tokens + output_tokens)
    save_usage_data(data)

# ==========================================
# 2. 🎨 UI 美化
# ==========================================
def inject_custom_css():
    tech_wave_bg = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 1440 320'>
      <path fill='none' stroke='%23C9CACA' stroke-width='1.5' stroke-opacity='0.5' d='M0,160L48,176C96,192,192,224,288,224C384,224,480,192,576,165.3C672,139,768,117,864,128C960,139,1056,181,1152,197.3C1248,213,1344,203,1392,197.3L1440,192' />
      <path fill='none' stroke='%231F323D' stroke-width='1' stroke-opacity='0.2' d='M0,224L48,213.3C96,203,192,181,288,181.3C384,181,480,203,576,218.7C672,235,768,245,864,229.3C960,213,1056,171,1152,149.3C1248,128,1344,128,1392,128L1440,128' />
    </svg>
    """
    tech_wave_bg = tech_wave_bg.replace('\n', '').strip()

    st.markdown(f"""
        <style>
        [data-testid="stSidebar"] {{
            background-color: rgba(180, 196, 63, 0.5);
            backdrop-filter: blur(10px);
            border-right: 1px solid rgba(180, 196, 63, 0.3);
        }}
        [data-testid="stSidebar"] .stMarkdown, 
        [data-testid="stSidebar"] h1, 
        [data-testid="stSidebar"] h2, 
        [data-testid="stSidebar"] h3, 
        [data-testid="stSidebar"] label,
        [data-testid="stSidebar"] .caption {{
            color: #1F323D !important;
        }}
        .stApp {{
            background-color: #ffffff;
            background-image: url("data:image/svg+xml;utf8,{tech_wave_bg}");
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
        }}
        .block-container {{
            background-color: rgba(255, 255, 255, 0.95);
            border-radius: 20px;
            padding: 3rem;
            margin-top: 2rem;
            box-shadow: 0 10px 30px rgba(31, 50, 61, 0.08);
            border: 1px solid rgba(201, 202, 202, 0.3);
        }}
        div.stButton > button:first-child {{
            background: linear-gradient(135deg, #1F323D 0%, #354A56 100%);
            color: white;
            font-size: 18px;
            font-weight: bold;
            border-radius: 8px;
            border: none;
            padding: 0.6rem 1rem;
            width: 100%;
            transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1);
            box-shadow: 0 4px 12px rgba(31, 50, 61, 0.2);
        }}
        div.stButton > button:first-child:hover {{
            transform: translateY(-2px);
            box-shadow: 0 8px 16px rgba(31, 50, 61, 0.3);
            background: linear-gradient(135deg, #2A4250, #4A6273);
        }}
        .info-card {{
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 12px;
            border-left: 6px solid #B4C43F;
            margin-bottom: 25px;
            color: #333;
            font-size: 1.05rem;
        }}
        .usage-metric-box {{
            border: 1px solid #1F323D;
            border-radius: 10px;
            padding: 12px;
            text-align: center;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
            transition: all 0.3s ease;
        }}
        .usage-metric-title {{ font-size: 0.9em; font-weight: 600; margin-bottom: 4px;}}
        .usage-metric-value {{ font-size: 1.4em; font-weight: 800; }}
        h1 {{ color: #1F323D; font-weight: 800; }}
        #MainMenu {{visibility: hidden;}}
        footer {{visibility: hidden;}}
        </style>
    """, unsafe_allow_html=True)

# ==========================================
# 3. 系統提示詞
# ==========================================
SYSTEM_INSTRUCTION = """
你是一位專業的行政秘書。請分析使用者提供的檔案（文件、錄音或圖片），並根據使用者的要求產出對應的 JSON 資料。
請嚴格遵守以下規則：

1. **Memo (指定格式)**：
   若任務是 Memo，請回傳 JSON 包含以下欄位。
   **重要：針對 'method', 'official', 'note' 等勾選欄位，請輸出「包含所有選項的完整字串」，並將判斷應勾選的項目符號改為「實心方塊 ■」，未選項目維持「空心方塊 □」。**
   {
       "time": "時間 (請完整填寫，如：113年12月25日 14:00)",
       "location": "地點",
       "method": "方式 (例如：'□電話 □活動 ■會議 □公文批示 □其他')",
       "official": "長官 (例如：'■部長 □次長 □主任秘書 □立法委員 □其他：')",
       "meeting_name": "會議名稱",
       "chair": "主席",
       "attendees": "出席人員",
       "related_dept": "相關部會",
       "guest_dept": "列席單位",
       "conclusions": ["結論1 (請以條列式呈現)", "結論2"],
       "action_items": ["辦理事項1 (請以條列式呈現)", "辦理事項2"],
       "note": "附言 (例如：'□請回電話 □請惠處 ■請參酌 □其他')",
       "filename_prefix": "建議檔名 (不含副檔名)"
   }

2. **簡易開會通知單 (指定格式)**：
   若任務是開會通知，請回傳 JSON 包含以下欄位：
   {
       "date": "發文日期 (例如: 113年12月25日)",
       "dept": "發文單位 (例如: 政策規劃組)",
       "reason": "開會事由",
       "full_time": "開會完整時間 (例如: 113年12月30日(星期二) 下午 4:00 - 5:00)",
       "location": "地點",
       "host": "主持人",
       "attendees": "出席人員 (若無資訊填寫 '詳如簽到表')",
       "note": "簡述/討論議題說明",
       "agenda_table": [ ["時間1", "主題1", "備註1"], ["時間2", "主題2", "備註2"] ],
       "filename_prefix": "建議檔名"
   }

3. **談參 (指定歸納邏輯)**：
   若任務是談參，請回傳 JSON 包含以下三個主要區塊：
   {
       "title": "談參主題",
       "background": ["背景說明點1", "背景說明點2"], 
       "discussion_points": [
           {"subtitle": "小標題 (5-10字)", "content": "詳細說明 (50-100字)"},
           {"subtitle": "小標題 (5-10字)", "content": "詳細說明 (50-100字)"}
       ],
       "unit_opinion": "單位意見與立場說明 (請整合為一段完整的發言內容)",
       "filename_prefix": "建議檔名"
   }
   **邏輯規則：**
   - **背景說明**：請歸納 1-2 點背景資訊。
   - **討論重點**：請提供 5-10 點。每點必須包含一個「5-10字的小標題」以及對應的內容。
   - **單位意見**：請基於單位立場，提出具體的發言建議或立場聲明。

4. **數據提取 (Excel)**：
   請回傳一個 List，包含多個 Dictionary，每個 Dictionary 代表一行數據。

5. **語言與翻譯強制規則**：
   - **所有輸出內容必須為「繁體中文 (Traditional Chinese, Taiwan)」**。
   - 若原始資料包含外文，請務必先將其**翻譯並潤飾**為通順的繁體中文。
"""

# ==========================================
# 4. Gemini API 分析函數
# ==========================================
def analyze_content_with_gemini(file_list, task_type, api_key, user_instruction=""):
    if not api_key:
        return {"error": "請先在側邊欄輸入 API Key"}
    if not file_list:
        return {"error": "請至少上傳一個檔案"}

    genai.configure(api_key=api_key)
    generation_config = {
        "temperature": 0.2, 
        "response_mime_type": "application/json"
    }

    model_priority_list = [
        "gemini-2.5-flash",      
        "gemini-3.0-flash",      
        "gemini-2.5-flash-lite"  
    ]
    
    content_parts = []
    base_prompt = f"你是一位專業行政秘書。請分析接下來提供的多份文件，並製作：{task_type}。請注意：若不同文件內容有衝突，請以「日期較新」或「使用者補充指令」為主。"
    content_parts.append(base_prompt)

    file_inventory = []

    for uploaded_file in file_list:
        file_name = uploaded_file.name
        file_inventory.append(file_name)
        file_bytes = uploaded_file.getvalue()
        mime_type = uploaded_file.type
        
        if file_name.lower().endswith('.m4a'):
             mime_type = 'audio/mp4'

        content_parts.append(f"\n=== 檔案開始：{file_name} ===\n")

        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = Document(BytesIO(file_bytes))
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        full_text.append(" | ".join(row_text))
                
                extracted_text = "\n".join(full_text)
                content_parts.append(extracted_text)
            except Exception as e:
                return {"error": f"檔案 {file_name} 讀取失敗: {str(e)}"}
        else:
            content_parts.append({
                "mime_type": mime_type,
                "data": file_bytes
            })
        
        content_parts.append(f"\n=== 檔案結束：{file_name} ===\n")

    final_instruction_block = f"""
    \n
    ---
    **資料清單**：{', '.join(file_inventory)}
    
    【重要：使用者特別補充指令】
    請在分析上述檔案時，優先遵守以下指示：
    {user_instruction if user_instruction else "無特別指令，請依照標準格式產出。"}
    
    **注意**：
    1. 若上述「使用者指令」與檔案內容有出入，請以「使用者指令」為準。
    2. 請務必輸出純 JSON 格式。
    ---
    """
    content_parts.append(final_instruction_block)

    status_container = st.status("🤖 AI 行政秘書正在多模態分析中...", expanded=True)
    last_error = ""

    for model_name in model_priority_list:
        try:
            status_container.write(f"正在呼叫模型：**{model_name}** ...")
            model = genai.GenerativeModel(
                model_name=model_name,
                generation_config=generation_config,
                system_instruction=SYSTEM_INSTRUCTION
            )
            response = model.generate_content(content_parts)
            
            if not response.text:
                raise ValueError("API 回傳空值")

            json_result = json.loads(response.text)
            
            if hasattr(response, 'usage_metadata'):
                usage = response.usage_metadata
                input_t = usage.prompt_token_count
                output_t = usage.candidates_token_count
                update_usage_count(model_name, input_t, output_t)
                json_result['_meta_info'] = {
                    "model": model_name,
                    "input_tokens": input_t,
                    "output_tokens": output_t,
                    "total_tokens": usage.total_token_count
                }

            status_container.update(label=f"✅ 分析完成！使用模型：{model_name}", state="complete", expanded=False)
            return json_result

        except Exception as e:
            error_msg = str(e)
            last_error = error_msg
            status_container.write(f"⚠️ {model_name} 發生錯誤: {error_msg}，切換備援...")
            continue

    status_container.update(label="❌ 所有模型皆失敗", state="error")
    return {"error": f"所有模型嘗試皆失敗。最後錯誤: {last_error}"}

# ==========================================
# 5. 檔案生成函數
# ==========================================
def set_chinese_font(run, font_name='標楷體', size_pt=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# --- Memo (模板模式 + 舊版備援) ---
def create_memo_docx_legacy(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.font.size = Pt(12)
    doc.add_paragraph("⚠️ (備援模式) 未偵測到 Template_Memo.docx，僅列出純文字內容。")
    doc.add_paragraph(f"時間：{data.get('time', '')}")
    doc.add_paragraph(f"內容：\n{data.get('conclusions', '')}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, "Legacy_Memo.docx"

def create_memo_docx(data):
    default_template_path = "Template_Memo.docx"
    if not os.path.exists(default_template_path):
        return create_memo_docx_legacy(data)
    try:
        doc = DocxTemplate(default_template_path)
        context = {
            'time': data.get('time', ''),
            'location': data.get('location', ''),
            'method': data.get('method', ''),     
            'official': data.get('official', ''), 
            'meeting_name': data.get('meeting_name', ''),
            'chair': data.get('chair', ''),
            'attendees': data.get('attendees', ''),
            'related_dept': data.get('related_dept', ''),
            'guest_dept': data.get('guest_dept', ''),
            'conclusions': data.get('conclusions', []), 
            'action_items': data.get('action_items', []),
            'note': data.get('note', ''),
            'filename_prefix': data.get('filename_prefix', 'Memo')
        }
        doc.render(context)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio, f"{context['filename_prefix']}.docx"
    except Exception as e:
        st.error(f"❌ Memo 模板生成失敗: {str(e)}")
        return create_memo_docx_legacy(data)

# --- 開會通知單 (模板模式 + 舊版備援) ---
def create_notice_docx_legacy(data):
    doc = Document()
    doc.add_paragraph("⚠️ 錯誤：找不到模板檔案，且已切換至備援模式。")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, "Legacy_Notice.docx"

def create_notice_docx(data, custom_template=None):
    default_template_path = "Template_Notice.docx" 
    doc = None
    if custom_template:
        doc = DocxTemplate(custom_template)
    elif os.path.exists(default_template_path):
        doc = DocxTemplate(default_template_path)
    else:
        return create_notice_docx_legacy(data)
    try:
        agenda_list = []
        if 'agenda_table' in data and isinstance(data['agenda_table'], list):
            for item in data['agenda_table']:
                col1 = str(item[0]) if len(item) > 0 else ""
                col2 = str(item[1]) if len(item) > 1 else ""
                col3 = str(item[2]) if len(item) > 2 else ""
                agenda_list.append({'col1': col1, 'col2': col2, 'col3': col3})

        context = {
            'date': data.get('date', ''),
            'dept': data.get('dept', ''),
            'reason': data.get('reason', ''),
            'full_time': data.get('full_time', ''),
            'location': data.get('location', ''),
            'host': data.get('host', ''),
            'attendees': data.get('attendees', ''),
            'summary': data.get('note', ''),
            'agenda_table': agenda_list, 
            'filename_prefix': data.get('filename_prefix', 'MeetingNotice')
        }
        doc.render(context)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio, f"{context['filename_prefix']}.docx"
    except Exception as e:
        st.error(f"❌ 模板生成失敗: {str(e)}")
        return create_notice_docx_legacy(data)

# --- 談參 (維持 Code 模式) ---
def create_talking_points_docx(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.font.size = Pt(12)
    
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(data.get('title', '談參資料'))
    r_title.bold = True
    set_chinese_font(r_title, size_pt=18)
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("-" * 30).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if data.get('background'):
        p_h1 = doc.add_paragraph()
        r_h1 = p_h1.add_run("一、背景說明")
        r_h1.bold = True
        set_chinese_font(r_h1, size_pt=14)
        for item in data['background']:
            p = doc.add_paragraph(style='List Bullet')
            set_chinese_font(p.add_run(item))

    if data.get('discussion_points'):
        p_h2 = doc.add_paragraph()
        r_h2 = p_h2.add_run("二、討論重點")
        r_h2.bold = True
        set_chinese_font(r_h2, size_pt=14)
        for item in data['discussion_points']:
            p = doc.add_paragraph(style='List Number')
            if 'subtitle' in item:
                r_sub = p.add_run(f"【{item['subtitle']}】")
                r_sub.bold = True
                set_chinese_font(r_sub)
            if 'content' in item:
                r_con = p.add_run(f"：{item['content']}")
                set_chinese_font(r_con)

    if data.get('unit_opinion'):
        p_h3 = doc.add_paragraph()
        r_h3 = p_h3.add_run("三、單位意見")
        r_h3.bold = True
        set_chinese_font(r_h3, size_pt=14)
        p_op = doc.add_paragraph()
        p_op.paragraph_format.first_line_indent = Pt(24)
        set_chinese_font(p_op.add_run(data['unit_opinion']))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, f"{data.get('filename_prefix', 'TalkingPoints')}.docx"

# --- Excel & Sheets ---
def create_excel(data_list):
    df = pd.DataFrame(data_list)
    bio = BytesIO()
    with pd.ExcelWriter(bio, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Data_Extraction')
    bio.seek(0)
    return bio, "Data_Extraction.xlsx"

def create_google_sheet(data, task_type, creds_dict, user_email=None):
    SCOPES = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
    try:
        creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
        client = gspread.authorize(creds)
        title = f"{data.get('filename_prefix', 'Export')}_{datetime.now().strftime('%m%d_%H%M')}"
        sh = client.create(title)
        ws = sh.sheet1
        rows_to_write = []
        
        if task_type == "數據提取 (Excel)":
             if isinstance(data, list) and len(data) > 0:
                 header = list(data[0].keys())
                 rows_to_write.append(header)
                 for item in data: rows_to_write.append(list(item.values()))
             else:
                 ws.update_acell('A1', '無數據')
        else:
            for k, v in data.items():
                if isinstance(v, list):
                    if not v: continue
                    if all(isinstance(x, str) for x in v):
                         rows_to_write.append([k, "\n".join(v)])
                    elif all(isinstance(x, dict) for x in v):
                         rows_to_write.append([k, "(詳如下表)"])
                         for sub_item in v:
                             rows_to_write.append(["", sub_item.get('subtitle',''), sub_item.get('content','')])
                else:
                    rows_to_write.append([k, str(v)])
        
        if rows_to_write: ws.update(rows_to_write)
        if user_email: sh.share(user_email, perm_type='user', role='writer')
        else: sh.share(None, perm_type='anyone', role='writer')
        return sh.url, "✅ 成功建立 Google Sheet！"
    except Exception as e:
        return None, f"❌ 錯誤: {str(e)}"

# ==========================================
# 6. Streamlit UI 主程式
# ==========================================
def main():
    inject_custom_css()

    with st.sidebar:
        st.title("⚙️ 設定面板")
        
        # 顯示模板狀態
        tpl_notice_exist = os.path.exists("Template_Notice.docx")
        tpl_memo_exist = os.path.exists("Template_Memo.docx")
        
        # -----------------------------------------------------
        # [修改] 用量統計與視覺化警示
        # -----------------------------------------------------
        st.markdown("### 📊 今日用量統計")
        usage_data = load_usage_data()
        target_models = ["gemini-2.5-flash", "gemini-3.0-flash", "gemini-2.5-flash-lite"]
        
        for m in target_models:
            count = usage_data["stats"].get(m, {}).get("count", 0)
            
            # 定義顏色邏輯
            if count >= 15:
                bg_color = "#D32F2F"   # 深紅
                text_color = "#FFFFFF" # 白字
                sub_text_color = "#EEEEEE" # 次要文字也反白
            elif count >= 10:
                bg_color = "#FBC02D"   # 黃色
                text_color = "#1F323D" # 深色字
                sub_text_color = "#1F323D"
            else:
                bg_color = "rgba(255, 255, 255, 0.6)" # 預設白
                text_color = "#1F323D"
                sub_text_color = "#1F323D"

            st.markdown(f"""
            <div class="usage-metric-box" style="margin-bottom: 8px; background-color: {bg_color};">
                <div class="usage-metric-title" style="color: {text_color};">{m}</div>
                <div class="usage-metric-value" style="color: {text_color};">
                    {count} <span style="font-size:0.5em;color: {sub_text_color};">次</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
        # -----------------------------------------------------

        st.markdown("---")
        
        with st.expander("☁️ Google Sheets 設定", expanded=False):
            uploaded_key = st.file_uploader("上傳 JSON Key", type=['json'], key="sheet_key")
            user_email = st.text_input("您的 Google Email (選填)")
            
        api_key = st.text_input("🔑 API Key", type="password", help="請輸入您的 Google Gemini API Key")
        
        st.subheader("📝 任務選擇")
        task_mode = st.radio(
            "請選擇輸出類型：",
            ("Memo (指定格式)", "簡易開會通知單 (指定格式)", "談參", "數據提取 (Excel)", "會議紀錄"),
            index=0
        )
        
        # 內建模板偵測與覆寫 UI
        custom_template_file = None
        if task_mode == "簡易開會通知單 (指定格式)":
            st.markdown("---")
            st.markdown("##### 📄 模板狀態")
            if tpl_notice_exist:
                st.caption("✅ 使用內建：Template_Notice.docx")
                if st.checkbox("手動上傳其他模板 (覆寫)"):
                    custom_template_file = st.file_uploader("上傳暫用模板 (.docx)", type=['docx'])
            else:
                st.warning("⚠️ 未偵測到內建模板")
                custom_template_file = st.file_uploader("請上傳模板 (.docx)", type=['docx'])
        
        if task_mode == "Memo (指定格式)":
            st.markdown("---")
            st.markdown("##### 📄 模板狀態")
            if tpl_memo_exist:
                st.caption("✅ 使用內建：Template_Memo.docx")
            else:
                st.warning("⚠️ 未偵測到內建模板 (Template_Memo.docx)")
                st.caption("請將模板檔案放入資料夾，否則將使用純文字模式")

        # 條件式補充指令
        user_instruction = ""
        if task_mode in ["談參", "數據提取 (Excel)", "Memo (指定格式)", "會議紀錄", "簡易開會通知單 (指定格式)"]:
            st.markdown("---")
            st.markdown(f"##### ✍️ 特別指示 (選填)")
            hint_text = "例如：請特別著重於... (此指令權重最高)"
            user_instruction = st.text_area("補充指令 (AI 將優先遵守)", placeholder=hint_text, height=100)

        st.caption("ADI Policy Planning AI Agent | Tech Wave Ed.")

    col1, col2 = st.columns([3, 1])
    with col1:
        st.title("🤖 數位產業署政策規劃組行政秘書")
        st.markdown("#### 自動化公文生成系統 | 支援多檔案、錄音與模板填充")
    with col2:
        st.markdown("")

    st.markdown('<div class="info-card">💡 系統提示：支援多檔案上傳。請在左側選擇任務與輸入指令，分析結果將自動優化為標準公文格式。</div>', unsafe_allow_html=True)

    with st.container(border=True):
        uploaded_files = st.file_uploader(
            "📂 拖放檔案到這裡或點擊上傳 (可多選)", 
            type=['docx', 'pdf', 'txt', 'wav', 'mp3', 'm4a', 'png', 'jpg', 'pptx'],
            accept_multiple_files=True
        )

    if uploaded_files:
        col_preview, col_action = st.columns([1, 2])
        with col_preview:
            st.info(f"📎 已上傳 {len(uploaded_files)} 個檔案")
            for f in uploaded_files:
                st.caption(f"- {f.name}")
        
        with col_action:
            if st.button("🚀 開始智慧分析"):
                if not api_key:
                    st.toast("⚠️ 請先在側邊欄輸入 API Key", icon="🔑")
                else:
                    result = analyze_content_with_gemini(uploaded_files, task_mode, api_key, user_instruction)
                    if "error" in result:
                        st.error(result["error"])
                    else:
                        st.session_state['result_data'] = result
                        if '_meta_info' in result:
                             st.session_state['meta_info'] = result.pop('_meta_info')
                        else:
                             st.session_state['meta_info'] = None
                        st.rerun()

    if 'result_data' in st.session_state and st.session_state['result_data']:
        result_data = st.session_state['result_data']
        meta_info = st.session_state.get('meta_info')
        
        st.divider()
        st.subheader("📊 分析結果")
        if meta_info:
            m_col1, m_col2, m_col3 = st.columns(3)
            m_col1.metric("使用模型", meta_info['model'])
            m_col2.metric("輸入 Token", f"{meta_info['input_tokens']:,}")
            m_col3.metric("輸出 Token", f"{meta_info['output_tokens']:,}")

        tab1, tab2, tab3 = st.tabs(["📥 下載產出", "🔍 原始資料 (JSON)", "📋 數據表格"])

        with tab1:
            st.success("文件已生成！請點擊下方按鈕下載。")
            if task_mode == "Memo (指定格式)":
                file_bio, file_name = create_memo_docx(result_data)
                st.download_button("📥 下載 Memo Word 檔", file_bio, file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            elif task_mode == "簡易開會通知單 (指定格式)":
                file_bio, file_name = create_notice_docx(result_data, custom_template_file)
                st.download_button("📥 下載 開會通知單 Word 檔", file_bio, file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            elif task_mode == "談參":
                file_bio, file_name = create_talking_points_docx(result_data)
                st.download_button("📥 下載 談參 Word 檔", file_bio, file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            elif task_mode == "數據提取 (Excel)":
                file_bio, file_name = create_excel(result_data)
                st.download_button("📥 下載 Excel 數據表", file_bio, file_name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            else:
                st.download_button("📥 下載文字檔 (.txt)", str(result_data), "result.txt", use_container_width=True)

            st.markdown("---")
            if st.button("📤 同步生成 Google Sheet", use_container_width=True):
                if uploaded_key is None:
                    st.error("⚠️ 請先在側邊欄上傳 Service Account JSON Key")
                else:
                    try:
                        stringio = BytesIO(uploaded_key.getvalue())
                        creds_dict = json.load(stringio)
                        with st.spinner("正在建立 Google Sheet..."):
                            sheet_url, msg = create_google_sheet(result_data, task_mode, creds_dict, user_email)
                        if sheet_url:
                            st.success(msg)
                            st.markdown(f"🔗 [點擊開啟 Google Sheet]({sheet_url})")
                        else:
                            st.error(msg)
                    except Exception as e:
                        st.error(f"認證檔案讀取錯誤: {e}")

        with tab2:
            st.json(result_data)

        with tab3:
            if task_mode == "簡易開會通知單 (指定格式)" and 'agenda_table' in result_data:
                st.dataframe(pd.DataFrame(result_data['agenda_table'], columns=['時間', '主題', '備註']), use_container_width=True)
            elif task_mode == "談參" and 'discussion_points' in result_data:
                st.dataframe(pd.DataFrame(result_data['discussion_points']), use_container_width=True)
            elif task_mode == "數據提取 (Excel)" and isinstance(result_data, list):
                st.dataframe(result_data, use_container_width=True)
            elif task_mode == "Memo (指定格式)" and 'action_items' in result_data:
                st.caption("辦理事項清單")
                st.dataframe(pd.DataFrame(result_data['action_items'], columns=['待辦事項']), use_container_width=True)
            else:
                st.info("此模式無預覽表格")

if __name__ == "__main__":
    main()
